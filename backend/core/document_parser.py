import re
from pathlib import Path
from typing import List, Dict

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

def clean_text(text: str) -> str:
    """
    Cleans raw text by removing non-printable characters and normalizing whitespace.
    """
    if not text:
        return ""
    # Normalize whitespaces to single spaces
    text = re.sub(r"\s+", " ", text)
    # Remove control characters but keep standard punctuation and letters
    text = re.sub(r"[^\x20-\x7E\n\t]", "", text)
    return text.strip()

def chunk_text(text: str, chunk_size: int = 600, overlap: int = 150) -> List[str]:
    """
    Splits text into overlapping chunks of a target character length.
    Ensures that text splits happen at word boundaries (spaces) to preserve meaning.
    """
    chunks = []
    if not text:
        return chunks
    
    text_len = len(text)
    start = 0
    
    while start < text_len:
        # Initial end position
        end = min(start + chunk_size, text_len)
        
        # If we are not at the end of the document, look back to find a word boundary
        if end < text_len:
            # Search for a space in the last 15% of the chunk to split cleanly
            lookback_limit = int(chunk_size * 0.15)
            last_space = text.rfind(" ", end - lookback_limit, end)
            if last_space != -1:
                end = last_space
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
            
        # If we reached the end of the document, we stop
        if end >= text_len:
            break
            
        # Move start position back by overlap
        start = end - overlap
        
        # Guard against zero progress or infinite loops
        if start >= end:
            start = end + 1
            
    return chunks

def extract_text_from_file(file_path: Path) -> str:
    """
    Extracts raw text from a document based on its file extension.
    Supports .txt, .md, and .pdf.
    """
    extension = file_path.suffix.lower()
    
    if extension in [".txt", ".md"]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
            
    elif extension == ".pdf":
        if pypdf is None:
            raise ImportError(
                "The 'pypdf' library is required to parse PDF files. "
                "Please install it using: pip install pypdf"
            )
            
        text_parts = []
        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                    
        return "\n".join(text_parts)
        
    elif extension == ".docx":
        if docx is None:
            raise ImportError(
                "The 'python-docx' library is required to parse DOCX files. "
                "Please install it using: pip install python-docx"
            )
            
        doc = docx.Document(file_path)
        text_parts = []
        
        # Extract paragraph text
        for paragraph in doc.paragraphs:
            val = paragraph.text.strip()
            if val:
                text_parts.append(val)
                
        # Extract table cells cleanly, converting rows to pipe-separated contexts
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    # De-duplicate adjacent identical cells (occurs with merged cells in python-docx)
                    if cell_text and (not row_texts or cell_text != row_texts[-1]):
                        row_texts.append(cell_text)
                if row_texts:
                    text_parts.append(" | ".join(row_texts))
                    
        return "\n".join(text_parts)
        
    else:
        raise ValueError(f"Unsupported file type: {extension}")

def process_document(file_path: Path, chunk_size: int = 600, overlap: int = 150) -> List[str]:
    """
    Overhauled: Structure-aware semantic chunker that preserves heading hierarchies and groups
    text logically by document sections. Prepend section path to each chunk to contextualize matches.
    """
    extension = file_path.suffix.lower()
    
    if extension not in [".docx", ".pdf", ".txt", ".md"]:
        raise ValueError(f"Unsupported file type: {extension}")
        
    breadcrumbs = ["", "", ""]  # [H1, H2, H3]
    chunks = []
    
    current_chunk_text = ""
    current_section_path = ""
    
    def flush_chunk():
        nonlocal current_chunk_text, current_section_path
        val = current_chunk_text.strip()
        if val:
            prefix = f"[Section: {current_section_path}] " if current_section_path else ""
            chunks.append(prefix + val)
            current_chunk_text = ""

    if extension == ".docx":
        if docx is None:
            raise ImportError("The 'python-docx' library is required to parse DOCX.")
        doc = docx.Document(file_path)
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            style_name = paragraph.style.name.lower() if paragraph.style else ""
            is_heading = False
            heading_level = 0
            
            if "heading 1" in style_name or "title" in style_name:
                is_heading = True
                heading_level = 1
            elif "heading 2" in style_name:
                is_heading = True
                heading_level = 2
            elif "heading 3" in style_name:
                is_heading = True
                heading_level = 3
            # Heuristic for bold inline runs acting as section headings
            elif len(text) < 80 and paragraph.runs and all(r.bold for r in paragraph.runs if r.text.strip()):
                is_heading = True
                heading_level = 2
                
            if is_heading:
                flush_chunk()
                lvl_idx = heading_level - 1
                if 0 <= lvl_idx < 3:
                    breadcrumbs[lvl_idx] = text
                    for j in range(lvl_idx + 1, 3):
                        breadcrumbs[j] = ""
                current_section_path = " > ".join([b for b in breadcrumbs if b])
            else:
                clean_p = clean_text(text)
                if not clean_p:
                    continue
                if len(current_chunk_text) + len(clean_p) > chunk_size:
                    flush_chunk()
                current_chunk_text += " " + clean_p
                
        flush_chunk()
        
        # Process tables
        for table in doc.tables:
            table_lines = []
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = clean_text(cell.text.strip())
                    if cell_text and (not row_texts or cell_text != row_texts[-1]):
                        row_texts.append(cell_text)
                if row_texts:
                    table_lines.append(" | ".join(row_texts))
            if table_lines:
                table_content = "\n".join(table_lines)
                prefix = f"[Section: {current_section_path} > Data Table] " if current_section_path else "[Section: Data Table] "
                chunks.append(prefix + table_content)
                
    elif extension in [".txt", ".md"]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            lines = f.readlines()
            
        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue
                
            is_heading = False
            heading_level = 0
            h_text = ""
            
            if line_str.startswith("# "):
                is_heading = True
                heading_level = 1
                h_text = line_str[2:].strip()
            elif line_str.startswith("## "):
                is_heading = True
                heading_level = 2
                h_text = line_str[3:].strip()
            elif line_str.startswith("### "):
                is_heading = True
                heading_level = 3
                h_text = line_str[4:].strip()
            elif len(line_str) < 60 and line_str.isupper() and not line_str.endswith((".", ",", ";")):
                is_heading = True
                heading_level = 2
                h_text = line_str
                
            if is_heading:
                flush_chunk()
                lvl_idx = heading_level - 1
                if 0 <= lvl_idx < 3:
                    breadcrumbs[lvl_idx] = h_text
                    for j in range(lvl_idx + 1, 3):
                        breadcrumbs[j] = ""
                current_section_path = " > ".join([b for b in breadcrumbs if b])
            else:
                clean_l = clean_text(line_str)
                if not clean_l:
                    continue
                if len(current_chunk_text) + len(clean_l) > chunk_size:
                    flush_chunk()
                current_chunk_text += " " + clean_l
                
        flush_chunk()
        
    elif extension == ".pdf":
        raw_text = extract_text_from_file(file_path)
        lines = raw_text.split("\n")
        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue
                
            is_heading = False
            heading_level = 0
            
            if len(line_str) < 60 and (line_str.isupper() or line_str.istitle()) and not line_str.endswith((".", ",", ";")):
                is_heading = True
                heading_level = 2
                h_text = line_str
                
            if is_heading:
                flush_chunk()
                breadcrumbs[1] = h_text
                breadcrumbs[2] = ""
                current_section_path = " > ".join([b for b in breadcrumbs if b])
            else:
                clean_l = clean_text(line_str)
                if not clean_l:
                    continue
                if len(current_chunk_text) + len(clean_l) > chunk_size:
                    flush_chunk()
                current_chunk_text += " " + clean_l
        flush_chunk()
        
    return [c.strip() for c in chunks if c.strip()]
